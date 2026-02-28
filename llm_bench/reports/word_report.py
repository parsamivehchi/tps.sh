"""Generate a polished, professional Word document report."""

import json
from io import BytesIO
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from rich.console import Console

from llm_bench.config import (
    REPORTS_DIR, RESULTS_DIR, SCORED_DIR,
    MODELS, CATEGORIES, CATEGORY_LABELS, QUALITY_WEIGHTS,
)
from llm_bench.reports.charts import (
    speed_chart, latency_chart, quality_chart, cost_chart,
    category_heatmap, combined_speed_quality, model_comparison_radar,
)

console = Console()

# ── Color constants ──
BRAND_BLUE = RGBColor(0x0F, 0x17, 0x2A)
ACCENT_BLUE = RGBColor(0x3B, 0x82, 0xF6)
DARK_TEXT = RGBColor(0x1E, 0x29, 0x3B)
MID_TEXT = RGBColor(0x47, 0x55, 0x69)
LIGHT_TEXT = RGBColor(0x94, 0xA3, 0xB8)
GREEN = RGBColor(0x10, 0xB9, 0x81)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
ROSE = RGBColor(0xEC, 0x48, 0x99)
TABLE_HEADER_BG = "0F172A"
TABLE_ALT_BG = "F8FAFC"


def _set_cell_bg(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def _style_paragraph(p, size=10, color=DARK_TEXT, bold=False, space_after=6):
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.bold = bold


def _add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level)
    for run in h.runs:
        run.font.color.rgb = BRAND_BLUE
    return h


def _add_body(doc, text, size=10.5, color=DARK_TEXT, space_after=8):
    p = doc.add_paragraph(text)
    _style_paragraph(p, size=size, color=color, space_after=space_after)
    return p


def _add_bullet(doc, text, level=0, size=10, color=DARK_TEXT):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    _style_paragraph(p, size=size, color=color, space_after=3)
    return p


def _add_callout(doc, text, icon=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{icon}  {text}" if icon else text)
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_BLUE
    run.font.italic = True
    return p


def _make_pro_table(doc, headers, rows, col_widths=None):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, TABLE_HEADER_BG)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            if ri % 2 == 0:
                _set_cell_bg(cell, TABLE_ALT_BG)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = DARK_TEXT

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def _add_chart(doc, buf, caption="", width=Inches(5.8)):
    if buf is None or buf.getbuffer().nbytes == 0:
        return
    doc.add_picture(buf, width=width)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_paragraph(p, size=9, color=LIGHT_TEXT, space_after=12)


def generate_word_report(run_id: str | None = None) -> Path:
    """Generate the polished LLM-Bench Word report."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LLM-Bench")
    run.font.size = Pt(36)
    run.font.color.rgb = BRAND_BLUE
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Local vs Cloud Model Benchmarking Report")
    run.font.size = Pt(16)
    run.font.color.rgb = MID_TEXT

    doc.add_paragraph()

    meta_text = f"Generated: {datetime.now().strftime('%B %d, %Y')}"
    if run_id:
        meta_file = RESULTS_DIR / run_id / "meta.json"
        if meta_file.exists():
            meta = json.loads(meta_file.read_text())
            meta_text += f"  |  Run ID: {run_id}"
            meta_text += f"  |  Tests: {meta.get('total_tests', '?')}"
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run(meta_text)
    run.font.size = Pt(10)
    run.font.color.rgb = LIGHT_TEXT

    doc.add_paragraph()
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run("\u2500" * 50)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        "A systematic comparison of 4 local open-source models (Ollama, M2 Max 32GB)\n"
        "against 3 cloud Claude models (Anthropic API) across 21 software engineering tasks."
    )
    run.font.size = Pt(10.5)
    run.font.color.rgb = MID_TEXT

    doc.add_page_break()

    # ── Table of Contents placeholder ──
    _add_styled_heading(doc, "Contents", 1)
    toc_items = [
        "1. Executive Summary",
        "2. Models Under Test",
        "3. Test Categories & Prompt Design",
        "4. Scoring Methodology",
        "5. Metrics Reference",
        "6. Results & Analysis",
        "7. Key Findings",
        "8. Interpreting Results",
        "9. Appendix: Bias Disclosure",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        _style_paragraph(p, size=11, color=ACCENT_BLUE, space_after=4)

    doc.add_page_break()

    # ═══════════════ Section 1: Executive Summary ═══════════════
    _add_styled_heading(doc, "1. Executive Summary", 1)
    _add_body(doc,
        "LLM-Bench is a benchmarking toolkit that systematically compares local open-source "
        "language models running on consumer Apple Silicon hardware against cloud-hosted Claude "
        "models from Anthropic. The benchmark measures three dimensions of model performance: "
        "speed (tokens per second, latency), quality (via LLM-as-judge scoring), and cost "
        "(API spend vs. zero-cost local inference)."
    )
    _add_body(doc,
        "The toolkit runs 21 carefully designed prompts across 7 software engineering categories "
        "against all 7 models, producing 147 individual test results. Each output is then scored "
        "by Claude Sonnet 4.6 on correctness, completeness, and clarity using a weighted rubric."
    )

    # Quick stats if we have data
    if run_id:
        results_file = RESULTS_DIR / run_id / "results.json"
        if results_file.exists():
            results = json.loads(results_file.read_text())
            success = [r for r in results if not r.get("error")]
            total_cost = sum(r.get("cost_usd", 0) for r in results)

            _add_callout(doc,
                f"This run completed {len(success)} tests across {len(set(r['model_name'] for r in success))} "
                f"models with a total cloud API cost of ${total_cost:.4f}."
            )

    # ═══════════════ Section 2: Models Under Test ═══════════════
    _add_styled_heading(doc, "2. Models Under Test", 1)
    _add_body(doc,
        "Seven models spanning two deployment paradigms: four local models running via Ollama "
        "on an Apple M2 Max with 32GB unified memory, and three cloud models accessed via the "
        "Anthropic Messages API."
    )

    _add_styled_heading(doc, "Local Models (Ollama)", 2)
    _add_body(doc,
        "Local models run entirely on-device with zero API cost. Performance is constrained by "
        "available hardware (M2 Max GPU cores, 32GB memory). Models are loaded one at a time "
        "to avoid memory pressure."
    )

    local_rows = []
    cloud_rows = []
    for m in MODELS:
        if m.provider == "ollama":
            local_rows.append([m.name, m.model_type, m.model_id, "Free"])
        else:
            cloud_rows.append([
                m.name, m.model_type, m.model_id,
                f"${m.cost_input:.0f} / ${m.cost_output:.0f}"
            ])

    _make_pro_table(doc,
        ["Model", "Type", "Model ID", "Cost"],
        local_rows,
        col_widths=[5, 3, 5, 2.5],
    )

    doc.add_paragraph()
    _add_styled_heading(doc, "Cloud Models (Anthropic)", 2)
    _add_body(doc,
        "Cloud models offer state-of-the-art quality with per-token pricing. "
        "All accessed via streaming API for accurate TTFT measurement."
    )

    _make_pro_table(doc,
        ["Model", "Type", "Model ID", "$/1M Tokens (In/Out)"],
        cloud_rows,
        col_widths=[4, 2.5, 4.5, 4.5],
    )

    # ═══════════════ Section 3: Test Categories ═══════════════
    doc.add_page_break()
    _add_styled_heading(doc, "3. Test Categories & Prompt Design", 1)
    _add_body(doc,
        "Prompts are organized into 7 categories covering the breadth of software engineering "
        "tasks where LLMs are commonly used. Each category contains 3 carefully crafted prompts "
        "that test different aspects of the skill, totaling 21 unique prompts per model."
    )

    cat_descriptions = {
        "code_generation": "Write functions, components, and full APIs from a specification. Tests creative coding ability.",
        "debugging_reasoning": "Find bugs in provided code, trace through logic, and apply chain-of-thought reasoning.",
        "short_quick": "One-liners, type conversions, and small utility functions. Tests speed on simple tasks.",
        "long_complex": "Multi-paragraph technical analysis, system design, and deep-dive explanations.",
        "refactoring": "Improve existing code: restructure, optimize queries, modernize syntax.",
        "explanation_teaching": "Tutorials, concept explainers, and educational content with examples.",
        "tool_calling": "Structured output, multi-tool orchestration, and agentic reasoning patterns.",
    }

    cat_rows = []
    for i, (cat_id, label) in enumerate(CATEGORY_LABELS.items(), 1):
        cat_rows.append([str(i), label, "3", cat_descriptions.get(cat_id, "")])

    _make_pro_table(doc,
        ["#", "Category", "Prompts", "Description"],
        cat_rows,
        col_widths=[1, 4, 1.5, 9],
    )

    _add_callout(doc, "Total per full run: 7 models x 21 prompts = 147 individual test executions")

    # ═══════════════ Section 4: Scoring Methodology ═══════════════
    _add_styled_heading(doc, "4. Scoring Methodology", 1)
    _add_body(doc,
        "Quality is assessed using an LLM-as-judge approach. Claude Sonnet 4.6 evaluates "
        "all 147 outputs (including outputs from other Claude models) against a structured rubric."
    )

    _add_styled_heading(doc, "Scoring Rubric", 2)
    score_rows = [
        ["Correctness", "40%", "Technical accuracy — are facts correct? Does code compile and work?"],
        ["Completeness", "35%", "Requirement coverage — are all parts of the prompt addressed?"],
        ["Clarity", "25%", "Organization and readability — is the response well-structured?"],
    ]
    _make_pro_table(doc,
        ["Dimension", "Weight", "What It Measures"],
        score_rows,
        col_widths=[3, 2, 10.5],
    )

    doc.add_paragraph()
    _add_body(doc,
        "Each dimension is scored on a 1\u201310 scale. The weighted composite score is computed as:"
    )
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run("Weighted Score = 0.40 \u00d7 Correctness + 0.35 \u00d7 Completeness + 0.25 \u00d7 Clarity")
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT_BLUE
    run.font.bold = True

    doc.add_paragraph()
    _add_styled_heading(doc, "Score Interpretation Guide", 2)
    interp_rows = [
        ["9.0 \u2013 10.0", "Exceptional", "Production-ready output, exceeds requirements"],
        ["7.0 \u2013 8.9", "Strong", "Solid response with minor room for improvement"],
        ["5.0 \u2013 6.9", "Adequate", "Usable but needs meaningful revision"],
        ["3.0 \u2013 4.9", "Weak", "Significant issues, partially addresses prompt"],
        ["1.0 \u2013 2.9", "Poor", "Major errors or largely misses the prompt"],
    ]
    _make_pro_table(doc,
        ["Range", "Rating", "Interpretation"],
        interp_rows,
        col_widths=[3, 3, 9.5],
    )

    # ═══════════════ Section 5: Metrics Reference ═══════════════
    doc.add_page_break()
    _add_styled_heading(doc, "5. Metrics Reference", 1)
    _add_body(doc, "Every test execution captures the following metrics:")

    metric_rows = [
        ["TTFT (ms)", "Time to First Token", "Latency from request to first response token. Measures perceived responsiveness."],
        ["Tokens/sec", "Generation Speed", "Output tokens generated per second. Higher = faster completion."],
        ["Total Time (ms)", "End-to-End Time", "Wall-clock time from request sent to full response received."],
        ["Input Tokens", "Prompt Size", "Number of tokens in the input prompt (varies by model tokenizer)."],
        ["Output Tokens", "Response Size", "Number of tokens in the model's response."],
        ["Quality (1\u201310)", "Weighted Score", "Composite quality score from LLM judge evaluation."],
        ["Cost ($)", "API Cost", "Calculated from token counts and per-model pricing. $0 for local models."],
        ["Memory (MB)", "VRAM Usage", "GPU memory consumed by the model (Ollama /api/ps). Local models only."],
    ]
    _make_pro_table(doc,
        ["Metric", "Full Name", "Description"],
        metric_rows,
        col_widths=[3, 3, 9.5],
    )

    # ═══════════════ Section 6: Results ═══════════════
    _add_styled_heading(doc, "6. Results & Analysis", 1)

    if run_id:
        _add_body(doc, "The following charts and tables present results from this benchmark run.")

        # Speed chart
        _add_styled_heading(doc, "6.1 Generation Speed", 2)
        _add_body(doc,
            "Average tokens per second across all prompts. Local models are typically faster "
            "for smaller outputs due to zero network latency, while cloud models may achieve "
            "higher sustained throughput."
        )
        _add_chart(doc, speed_chart(run_id), "Figure 1: Average generation speed (tokens/sec) by model")

        # Latency chart
        _add_styled_heading(doc, "6.2 Latency (Time to First Token)", 2)
        _add_body(doc,
            "Average TTFT measures how quickly a model begins responding. Local models avoid "
            "network round-trips but may have cold-start delays when loading into VRAM."
        )
        _add_chart(doc, latency_chart(run_id), "Figure 2: Average TTFT (ms) by model")

        # Quality charts
        buf = quality_chart(run_id)
        if buf:
            doc.add_page_break()
            _add_styled_heading(doc, "6.3 Quality Scores", 2)
            _add_body(doc,
                "Breakdown of quality sub-scores by model. Higher scores indicate better "
                "performance on that dimension."
            )
            _add_chart(doc, buf, "Figure 3: Quality sub-scores (correctness, completeness, clarity)")

        # Radar
        buf = model_comparison_radar(run_id)
        if buf:
            _add_chart(doc, buf, "Figure 4: Quality radar comparison across all models")

        # Heatmap
        buf = category_heatmap(run_id)
        if buf:
            _add_styled_heading(doc, "6.4 Category Performance", 2)
            _add_body(doc,
                "Performance varies by task type. Some models excel at code generation but struggle "
                "with long-form analysis, or vice versa."
            )
            _add_chart(doc, buf, "Figure 5: Quality heatmap (category x model)")

        # Speed vs Quality scatter
        buf = combined_speed_quality(run_id)
        if buf:
            _add_styled_heading(doc, "6.5 Speed vs Quality Trade-off", 2)
            _add_body(doc,
                "The ideal model appears in the upper-right (fast and high quality). "
                "Bubble size reflects average output length."
            )
            _add_chart(doc, buf, "Figure 6: Speed vs quality trade-off")

        # Cost chart
        buf = cost_chart(run_id)
        if buf:
            _add_styled_heading(doc, "6.6 API Costs", 2)
            _add_body(doc,
                "Total API spend for cloud models in this benchmark run. Local models incur "
                "zero marginal cost (electricity only)."
            )
            _add_chart(doc, buf, "Figure 7: Total API cost per cloud model")

        # Detailed results table
        scored_file = SCORED_DIR / run_id / "scored_results.json"
        if scored_file.exists():
            doc.add_page_break()
            _add_styled_heading(doc, "6.7 Full Results Table", 2)
            scored = json.loads(scored_file.read_text())

            detail_rows = []
            for r in sorted(scored, key=lambda x: (x.get("model_name", ""), x.get("prompt_id", ""))):
                scores = r.get("scores") or {}
                q_str = f"{scores['weighted']:.1f}" if isinstance(scores, dict) and scores.get("weighted") else "-"
                bias = " *" if r.get("self_bias_flag") else ""
                cost = f"${r.get('cost_usd', 0):.4f}" if r.get("cost_usd", 0) > 0 else "free"
                detail_rows.append([
                    r.get("model_name", ""),
                    r.get("prompt_id", ""),
                    f"{r.get('ttft_ms', 0):.0f}",
                    f"{r.get('tokens_per_sec', 0):.1f}",
                    str(r.get("output_tokens", 0)),
                    q_str + bias,
                    cost,
                ])

            _make_pro_table(doc,
                ["Model", "Prompt", "TTFT (ms)", "TPS", "Tokens", "Quality", "Cost"],
                detail_rows,
                col_widths=[3.5, 3.5, 2, 1.5, 1.5, 1.5, 2],
            )
    else:
        _add_callout(doc, "Run the benchmark first, then regenerate this report with results data.")
        _add_body(doc, "Charts and data tables will appear here after running:")
        _add_bullet(doc, "python -m llm_bench run")
        _add_bullet(doc, "python -m llm_bench judge <run_id>")
        _add_bullet(doc, "python -m llm_bench report <run_id>")

    # ═══════════════ Section 7: Key Findings ═══════════════
    doc.add_page_break()
    _add_styled_heading(doc, "7. Key Findings", 1)

    if run_id:
        analysis_file = SCORED_DIR / run_id / "analysis.json"
        if analysis_file.exists():
            analysis = json.loads(analysis_file.read_text())
            rankings = analysis.get("model_rankings", [])
            if rankings:
                fastest = min(rankings, key=lambda x: x.get("ttft_ms", float("inf")))
                best_tps = max(rankings, key=lambda x: x.get("tokens_per_sec", 0))
                _add_bullet(doc, f"Fastest TTFT: {fastest['model_name']} at {fastest['ttft_ms']:.0f}ms average")
                _add_bullet(doc, f"Highest throughput: {best_tps['model_name']} at {best_tps['tokens_per_sec']:.1f} tok/s")

                if analysis.get("has_quality_scores"):
                    best_q = max(rankings, key=lambda x: x.get("score_weighted", 0))
                    _add_bullet(doc, f"Best quality: {best_q['model_name']} at {best_q['score_weighted']:.2f}/10 weighted")

                paid = [r for r in rankings if r.get("cost_usd", 0) > 0]
                if paid:
                    cheapest = min(paid, key=lambda x: x["cost_usd"])
                    _add_bullet(doc, f"Most affordable cloud model: {cheapest['model_name']} at ${cheapest['cost_usd']:.4f} total")
        else:
            _add_body(doc, "Run analysis first to populate findings: python -m llm_bench analyze <run_id>")
    else:
        _add_body(doc, "[Findings will be generated from benchmark results]")

    # ═══════════════ Section 8: How to Interpret ═══════════════
    _add_styled_heading(doc, "8. Interpreting Results", 1)
    _add_body(doc,
        "The right model depends on your use case. Here are guidelines for different scenarios:"
    )

    guide_rows = [
        ["Batch code review", "Throughput (TPS)", "Local models — zero cost at scale"],
        ["Interactive chat", "Latency (TTFT)", "Lowest TTFT model — faster perceived response"],
        ["Critical production code", "Quality score", "Highest-quality model — usually cloud"],
        ["Quick prototyping", "Speed + adequate quality", "Fast local model with score > 6"],
        ["Cost-sensitive startup", "Quality per dollar", "Best cloud value or capable local model"],
    ]
    _make_pro_table(doc,
        ["Scenario", "Priority Metric", "Recommendation"],
        guide_rows,
        col_widths=[4, 3.5, 8],
    )

    # ═══════════════ Section 9: Bias Disclosure ═══════════════
    doc.add_paragraph()
    _add_styled_heading(doc, "9. Appendix: Bias Disclosure", 1)
    _add_body(doc,
        "Claude Sonnet 4.6 serves as the quality judge for all 147 outputs, including outputs "
        "from Claude Haiku 4.5, Claude Sonnet 4.6 itself, and Claude Opus 4.6. This creates "
        "an inherent self-evaluation bias — a model from the same family is scoring its own "
        "siblings and itself.",
        size=10,
    )
    _add_body(doc, "Mitigations applied:", size=10)
    _add_bullet(doc, "All Claude model scores are flagged with a bias indicator (*) in tables and the dashboard")
    _add_bullet(doc, "The judge uses a structured rubric to reduce subjective drift")
    _add_bullet(doc, "Scores are stored in editable JSON — users can manually override any score")
    _add_bullet(doc, "We recommend cross-validating with human review for high-stakes decisions")

    _add_body(doc,
        "Despite this limitation, Claude Sonnet was chosen as the judge for its strong "
        "cost/quality trade-off. Using an external judge (e.g., GPT-4) would eliminate family "
        "bias but introduce cross-vendor bias instead.",
        size=10, color=MID_TEXT,
    )

    # ── Footer ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by LLM-Bench — github.com/llm-bench")
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_TEXT

    # ── Save ──
    output_path = REPORTS_DIR / f"llm_bench_report{'_' + run_id if run_id else ''}.docx"
    doc.save(str(output_path))
    console.print(f"[green]Word report saved: {output_path}[/]")
    return output_path
